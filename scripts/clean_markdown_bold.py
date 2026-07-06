# -*- coding: utf-8 -*-
"""清理 DOCX 里残留的 markdown 加粗标记 `**xxx**`，转成真正的 DOCX 加粗。

用法:
    python clean_markdown_bold.py <docx_path>

前置:
    pip install python-docx

行为:
    - 找到所有 '**xxx**' 模式
    - 拆成三段 run：前文(普通) + xxx(加粗) + 后文(普通)
    - 保留原段落样式、对齐、其他格式
    - 自动备份原文件为 <docx_path>.bak
    - 如果 WPS 占用文件，提示用户先关闭

原理:
    python-docx 写 `p.add_run('**bold**')` 时，把整段当普通文本，
    星号原样进 DOCX。正确做法是分 run + run.bold=True。
"""
import sys
import os
import re
import shutil

# 避免 hermes 注入的 PIL 冲突
sys.path = [p for p in sys.path if 'hermes' not in p]

from docx import Document
from docx.oxml.ns import qn


def clean_markdown_bold(path: str) -> int:
    """清理 DOCX 里的 **xxx** 标记，返回清理的总数。"""
    if not os.path.exists(path):
        raise FileNotFoundError(f"DOCX 不存在: {path}")

    # 备份
    backup = path + '.bak'
    if not os.path.exists(backup):
        shutil.copy(path, backup)
        print(f"[backup] {backup}")

    doc = Document(path)
    total = 0

    for i, p in enumerate(doc.paragraphs):
        full_text = ''.join(r.text for r in p.runs)
        if '**' not in full_text:
            continue

        matches = list(re.finditer(r'\*\*([^*]+?)\*\*', full_text))
        if not matches:
            continue

        # 基础格式
        base_run = p.runs[0] if p.runs else None

        def make_run(text, bold=False):
            r = p.add_run(text)
            if base_run is not None:
                if base_run.font.name:
                    r.font.name = base_run.font.name
                    rPr = r._element.get_or_add_rPr()
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = rPr.makeelement(qn('w:rFonts'), {})
                        rPr.append(rFonts)
                    rFonts.set(qn('w:eastAsia'), base_run.font.name)
                if base_run.font.size:
                    r.font.size = base_run.font.size
                try:
                    if base_run.font.color and base_run.font.color.rgb:
                        r.font.color.rgb = base_run.font.color.rgb
                except (AttributeError, TypeError):
                    pass
            r.bold = bold
            return r

        # 清空原 run
        for r in list(p.runs):
            r._element.getparent().remove(r._element)

        # 重建
        cursor = 0
        for m in matches:
            start, end = m.span()
            if start > cursor:
                make_run(full_text[cursor:start], bold=False)
            make_run(m.group(1), bold=True)
            cursor = end
        if cursor < len(full_text):
            make_run(full_text[cursor:], bold=False)

        total += len(matches)
        print(f"  para#{i}: {len(matches)} 处加粗")

    try:
        doc.save(path)
        print(f"\n[ok] 共清理 {total} 处 ** 标记")
        print(f"[saved] {path}")
    except PermissionError as e:
        print(f"\n[ERROR] 保存失败: {e}", file=sys.stderr)
        print("→ 可能是 WPS 占用了文件。请关掉 WPS 窗口后重试。", file=sys.stderr)
        print(f"→ 备份在 {backup}，可手动恢复。", file=sys.stderr)
        raise

    return total


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python clean_markdown_bold.py <docx_path>")
        sys.exit(1)
    clean_markdown_bold(sys.argv[1])

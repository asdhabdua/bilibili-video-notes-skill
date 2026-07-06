# -*- coding: utf-8 -*-
"""
王道考研计算机网络 DOCX 笔记 v2 通用模板（严格按 SKILL.md 规范）

使用方法：
1. 复制此文件到工作目录（如 gen_pxx_v2.py）
2. 修改顶部配置区：TITLE / SOURCE / OUTPUT_PATH / FRAMES / SUMMARY / SECTIONS
3. 运行 python gen_pxx_v2.py

支持的内容块（SECTIONS 元素）：
  ('h1', '一、xxx')                    顶级大标题（一、二、三...）
  ('h2', '1. xxx')                      二级标题
  ('h3', '① xxx')                      三级标题
  ('body', '正文...')                  普通正文
  ('red', '关键术语')                 红色加粗正文
  ('bold', '重要标题')               黑色加粗正文
  ('bullet', [(text, bold, red), ...])  项目符号
  ('table', headers, rows)              表格
  ('img', '0001', 'caption')            嵌入图片（宽度占满页面）
  ('formula', 'C = λF')                 居中公式
  ('why', '解释...')                    WHY 解释（要完整：问题→原理→结论）
  ('tips', '做题要点...')             做题要点

SKILL.md 约束：
- 页边距：上下左右各 1.27 cm
- 图片宽度占满页面宽度（卡着页边距）
- 级别：第三章 → Heading 1；2.2 → Heading 2；2.2.1 → Heading 3
- 段落正文 Pt 13，行距 1.4，space_after = Pt(2)
- H1 标题 Pt 17，H2 标题 Pt 15，H3 标题 Pt 14
- 红色加粗 Pt 14，公式 Pt 15，图注/表格 Pt 11
- run.bold = True 真加粗（禁止 **xxx** 字面星号）
- 红色 RGBColor(192, 0, 0)

截图与字幕融会贯通原则（最重要）：
- 截图的价值是补齐字幕没讲到的考点，不是把图里所有文字抄进正文
- 字幕有的考点 → 用字幕写，简洁干练
- 字幕没有但图里有的考点 → 从图中补充
- 图里有但不是考点的 → 不写（如各层材料名、接口外观、线芯径）
- 判断标准：这个信息考研会考吗？不会考就不写进正文
- 不写"思维导图展示了""图中标注了""老师用红笔圈出"等描述性语句
- WHY 解释要完整（包含原因+原理+结论），不要砍成一句话
  正确示范: "为什么要加屏蔽层？外界环境中的电磁波会干扰线对上的数据信号。
  屏蔽层能把外部电场导走，减少进入导线的噪声功率，从而提高信噪比。
  根据香农定理，信噪比越高，极限数据传输速率越高。"
  错误示范: "屏蔽层把外部电磁干扰导走，降低噪声功率，提高信噪比，从而提高极限数据传输速率。"
  区别: 正确版先说"为什么需要"(问题场景)，再说"怎么解决"(原理)，最后说"结果"(香农定理)
"""
import os
import sys
import re

sys.path = [p for p in sys.path if 'hermes' not in p]

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

# =============================================================================
# === 每次需修改的配置区 ===
# =============================================================================

TITLE = '章节标题'
SOURCE = '视频来源说明'
OUTPUT_PATH = r'./notes/X.X_章节标题_v1.docx'
# FRAMES_DIR 必须指向当前视频独立的 selected/ 目录，禁止多视频共用
FRAMES_DIR = r'./frames/pXX/selected'

FRAMES = {
    # '0001': '图注文字',
}

SUMMARY = [
    # '要点1',
    # '要点2',
]

# =============================================================================
# === 以下是通用样式函数，一般不用改 ===
# =============================================================================

PAGE_MARGIN = Cm(1.27)
BODY_SIZE = 13
RED_COLOR = RGBColor(192, 0, 0)
H1_SIZE = 17
H2_SIZE = 15
H3_SIZE = 14
RED_SIZE = 14
FORMULA_SIZE = 15
LINE_SPACING_BODY = 1.4
SPACE_AFTER_P = Pt(2)

def set_run_font(run, name='宋体', size=BODY_SIZE, bold=False, red=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    if red:
        run.font.color.rgb = RED_COLOR

def _add_runs_with_sub(p, text, name='宋体', size=BODY_SIZE, bold=False, red=False):
    """添加文本 run，自动处理 _X 下标。下标 run 用 font.subscript 缩小+降低。"""
    parts = re.split(r'_([a-zA-Z0-9]+)', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        set_run_font(run, name=name, size=size, bold=bold, red=red)
        if i % 2 == 1:
            run.font.subscript = True

def setup_section(doc):
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = PAGE_MARGIN
    return s

def add_heading_styled(doc, text, level, size):
    style_name = f'Heading {level}'
    try:
        style = doc.styles[style_name]
    except KeyError:
        style = doc.styles.add_style(style_name, 1)
    p = doc.add_paragraph()
    p.style = style
    p.paragraph_format.space_before = Pt(8 if level == 1 else 6 if level == 2 else 4)
    p.paragraph_format.space_after = Pt(4 if level == 1 else 2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.2
    if p.runs:
        p._element.remove(p.runs[0]._element)
    _add_runs_with_sub(p, text, name='黑体', size=size, bold=True)
    return p

def add_h1(doc, text): return add_heading_styled(doc, text, 1, H1_SIZE)
def add_h2(doc, text): return add_heading_styled(doc, text, 2, H2_SIZE)
def add_h3(doc, text): return add_heading_styled(doc, text, 3, H3_SIZE)

def add_body(doc, text, bold=False, red=False, size=BODY_SIZE):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = LINE_SPACING_BODY
    _add_runs_with_sub(p, text, size=size, bold=bold, red=red)
    return p

def add_bullet(doc, parts, size=BODY_SIZE):
    """每个 parts 条目独立成一个 bullet 段落。"""
    for text, is_bold, is_red in parts:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = LINE_SPACING_BODY
        _add_runs_with_sub(p, text, size=size, bold=is_bold, red=is_red)
    return doc

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(h), name='黑体', size=BODY_SIZE, bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            p = table.rows[r_idx + 1].cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.add_run(cell_text), size=BODY_SIZE - 1)
    doc.add_paragraph()
    return table

def add_img(doc, frame_id, caption, img_width):
    path = os.path.join(FRAMES_DIR, f'frame_{frame_id}.jpg')
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=img_width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(2)
    set_run_font(cap.add_run(f'图：{caption}'), size=11)
    return p

def add_formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run(text), name='Times New Roman', size=FORMULA_SIZE, bold=True, red=True)
    return p

def add_why(doc, text):
    add_body(doc, 'WHY 解释：', bold=True, red=True, size=RED_SIZE)
    add_body(doc, text)

def add_tips(doc, text):
    add_body(doc, '做题要点', bold=True, red=True, size=RED_SIZE)
    add_bullet(doc, [(text, False, False)])

def main():
    doc = Document()
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    section = setup_section(doc)
    img_width = section.page_width - section.left_margin - section.right_margin

    if re.match(r'^第[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\u767e]+\u7ae0', TITLE):
        add_heading_styled(doc, TITLE, 1, H1_SIZE + 3)
    elif re.match(r'^\d+\.\d+\.\d+', TITLE):
        add_heading_styled(doc, TITLE, 3, H2_SIZE)
    elif re.match(r'^\d+\.\d+', TITLE):
        add_heading_styled(doc, TITLE, 2, H1_SIZE)
    else:
        add_heading_styled(doc, TITLE, 1, H1_SIZE + 3)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.paragraph_format.space_after = Pt(4)
    set_run_font(t2.add_run(f'来源：{SOURCE}'), size=11)
    doc.add_paragraph()

    # === SECTIONS 内容区（每次需修改） ===
    SECTIONS = [
        # 在这里填写正文内容
        # 参见 SKILL.md 中的排版规范和截图融会贯通原则
    ]

    for item in SECTIONS:
        typ = item[0]
        if typ == 'h1': add_h1(doc, item[1])
        elif typ == 'h2': add_h2(doc, item[1])
        elif typ == 'h3': add_h3(doc, item[1])
        elif typ == 'body': add_body(doc, item[1])
        elif typ == 'red': add_body(doc, item[1], bold=True, red=True, size=RED_SIZE)
        elif typ == 'bold': add_body(doc, item[1], bold=True)
        elif typ == 'bullet': add_bullet(doc, item[1])
        elif typ == 'table': add_table(doc, item[1], item[2])
        elif typ == 'img': add_img(doc, item[1], item[2], img_width)
        elif typ == 'formula': add_formula(doc, item[1])
        elif typ == 'why': add_why(doc, item[1])
        elif typ == 'tips': add_tips(doc, item[1])

    add_h1(doc, '本节要点总结')
    for s in SUMMARY:
        add_bullet(doc, [(s, False, False)])

    doc.core_properties.title = TITLE
    doc.core_properties.subject = SOURCE
    doc.core_properties.author = 'Hermes Agent'
    doc.save(OUTPUT_PATH)

    import zipfile
    with zipfile.ZipFile(OUTPUT_PATH) as z:
        with z.open('word/document.xml') as f:
            xml = f.read().decode('utf-8')
    n = len(re.findall(r'\*\*[^*]+\*\*', xml))
    print(f'Saved: {OUTPUT_PATH}\nSize: {os.path.getsize(OUTPUT_PATH)/1024:.1f} KB')
    print('No markdown residue.' if n == 0 else f'WARNING: {n} residues!')

if __name__ == '__main__':
    main()